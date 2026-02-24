"""
Script to create the pre-prepared COA Word template for docxtpl.

Run this script once to generate the template file:
    python -m modules.create_template
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "coa_template.docx")


def create_template():
    """Create the COA template docx file with Jinja2 placeholders for docxtpl."""
    os.makedirs(TEMPLATE_DIR, exist_ok=True)

    doc = Document()

    # Configure default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # Configure page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.8)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("{{ title }}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("{{ subtitle }}")
    sub_run.font.size = Pt(12)
    sub_run.font.name = "Times New Roman"
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    # Spacer
    doc.add_paragraph()

    # Metadata section
    meta_fields = [
        ("Исходный файл:", "{{ original_filename }}"),
        ("Дата перевода:", "{{ translation_date }}"),
        ("Модель перевода:", "{{ model_used }}"),
        ("Метод извлечения:", "{{ extraction_method }}"),
    ]

    for label, placeholder in meta_fields:
        p = doc.add_paragraph()
        label_run = p.add_run(label + " ")
        label_run.bold = True
        label_run.font.size = Pt(9)
        label_run.font.name = "Times New Roman"
        label_run.font.color.rgb = RGBColor(80, 80, 80)
        val_run = p.add_run(placeholder)
        val_run.font.size = Pt(9)
        val_run.font.name = "Times New Roman"
        val_run.font.color.rgb = RGBColor(80, 80, 80)

    # Divider
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = divider.add_run("─" * 60)
    div_run.font.size = Pt(8)
    div_run.font.color.rgb = RGBColor(180, 180, 180)

    # Main content placeholder
    content_para = doc.add_paragraph()
    content_run = content_para.add_run("{{ translated_content }}")
    content_run.font.name = "Times New Roman"
    content_run.font.size = Pt(11)

    # Bottom divider
    doc.add_paragraph()
    divider2 = doc.add_paragraph()
    divider2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run2 = divider2.add_run("─" * 60)
    div_run2.font.size = Pt(8)
    div_run2.font.color.rgb = RGBColor(180, 180, 180)

    # Disclaimer
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_run = disclaimer.add_run(
        "Данный документ является переводом оригинального Сертификата анализа.\n"
        "Перевод выполнен с использованием искусственного интеллекта с применением\n"
        "фармацевтического глоссария. Рекомендуется верификация специалистом."
    )
    disc_run.font.size = Pt(8)
    disc_run.font.name = "Times New Roman"
    disc_run.font.color.rgb = RGBColor(140, 140, 140)
    disc_run.italic = True

    doc.save(TEMPLATE_PATH)
    print(f"Template created at: {TEMPLATE_PATH}")


if __name__ == "__main__":
    create_template()
